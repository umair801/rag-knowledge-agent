import uuid
import threading
from fastapi import APIRouter, UploadFile, File, HTTPException, Query
from fastapi.responses import FileResponse
import tempfile
import os

from app.models.schemas import (
    TranslationRequest, TranslationResult,
    TranslationJob, JobStatusResponse
)
from app.translation.translator import translate_text, SUPPORTED_LANGUAGES
from app.translation.splitter import split_into_chapters
from app.translation.job_store import create_job, get_job, update_job
from app.ingestion.loaders import load_document

router = APIRouter()


# ---------------------------------------------------------------------------
# Health
# ---------------------------------------------------------------------------

@router.get("/health")
def health():
    return {"status": "ok", "service": "AI-Powered Book Translation Tool"}


# ---------------------------------------------------------------------------
# Languages
# ---------------------------------------------------------------------------

@router.get("/api/v1/translate/languages")
def list_languages():
    return {"supported_languages": SUPPORTED_LANGUAGES}


# ---------------------------------------------------------------------------
# Translate raw text block
# ---------------------------------------------------------------------------

@router.post("/api/v1/translate/text", response_model=TranslationResult)
def translate_text_endpoint(request: TranslationRequest):
    try:
        result = translate_text(request)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Translate uploaded file (PDF / DOCX / TXT)
# ---------------------------------------------------------------------------

@router.post("/api/v1/translate/file")
def translate_file_endpoint(
    file: UploadFile = File(...),
    target_language: str = Query(..., description="Target language name"),
    style: str = Query("literary", description="literary | formal | casual"),
):
    # Save upload to temp file
    suffix = os.path.splitext(file.filename)[-1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file.file.read())
        tmp_path = tmp.name

    job_id = str(uuid.uuid4())
    job = TranslationJob(
        job_id=job_id,
        status="pending",
        target_language=target_language,
        style=style,
    )
    create_job(job)

    # Run translation in background thread
    thread = threading.Thread(
        target=_run_file_translation,
        args=(job_id, tmp_path, target_language, style),
        daemon=True,
    )
    thread.start()

    return {
        "job_id": job_id,
        "status": "pending",
        "message": "Translation started. Poll /api/v1/jobs/{job_id} for progress.",
    }


def _run_file_translation(job_id: str, file_path: str, target_language: str, style: str):
    job = get_job(job_id)
    try:
        job.status = "processing"
        update_job(job)

        # Extract text (loader returns tuple: text, source_type, metadata)
        raw_text, _, _ = load_document(file_path)

        # Split into chapters
        chapters = split_into_chapters(raw_text)
        job.chapter_count = len(chapters)
        update_job(job)

        translated_chapters = []
        total_tokens = 0

        for chapter in chapters:
            req = TranslationRequest(
                source_text=chapter.content,
                target_language=target_language,
                style=style,
            )
            result = translate_text(req)
            total_tokens += result.token_usage.get("total_tokens", 0)
            translated_chapters.append({
                "chapter_number": chapter.chapter_number,
                "title": chapter.title,
                "translated_text": result.translated_text,
            })
            job.completed_chapters += 1
            job.total_tokens = total_tokens
            update_job(job)

        # Build output DOCX
        output_path = _build_docx(translated_chapters, target_language, job_id)

        job.status = "completed"
        job.total_tokens = total_tokens
        # Store output path in error field temporarily (repurposed as output_path)
        job.error = output_path
        update_job(job)

    except Exception as e:
        job.status = "failed"
        job.error = str(e)
        update_job(job)
    finally:
        os.unlink(file_path)


def _sanitize(text: str) -> str:
    """Remove characters that are invalid in XML/DOCX."""
    import re
    # Remove NULL bytes and non-XML-compatible control characters
    # Keep: tab (\x09), newline (\x0A), carriage return (\x0D)
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    return text


def _build_docx(translated_chapters: list, target_language: str, job_id: str) -> str:
    from docx import Document
    doc = Document()
    doc.add_heading(f"Translated Book — {target_language}", level=0)

    for ch in translated_chapters:
        title = _sanitize(ch["title"] or f"Chapter {ch['chapter_number']}")
        doc.add_heading(title, level=1)
        doc.add_paragraph(_sanitize(ch["translated_text"]))
        doc.add_page_break()

    output_path = os.path.join(tempfile.gettempdir(), f"translation_{job_id}.docx")
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Job status polling
# ---------------------------------------------------------------------------

@router.get("/api/v1/jobs/{job_id}", response_model=JobStatusResponse)
def get_job_status(job_id: str):
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    progress = 0.0
    if job.chapter_count > 0:
        progress = round((job.completed_chapters / job.chapter_count) * 100, 1)

    return JobStatusResponse(
        job_id=job.job_id,
        status=job.status,
        chapter_count=job.chapter_count,
        completed_chapters=job.completed_chapters,
        total_tokens=job.total_tokens,
        progress_percent=progress,
        error=job.error if job.status == "failed" else None,
    )


# ---------------------------------------------------------------------------
# Download completed translation
# ---------------------------------------------------------------------------

@router.get("/api/v1/jobs/{job_id}/download")
def download_translation(job_id: str):
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.status != "completed":
        raise HTTPException(status_code=400, detail=f"Job is not completed. Status: {job.status}")

    output_path = job.error  # repurposed field storing the output path
    if not output_path or not os.path.exists(output_path):
        raise HTTPException(status_code=404, detail="Output file not found")

    return FileResponse(
        path=output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"translation_{job_id}.docx",
    )


# ---------------------------------------------------------------------------
# All jobs listing
# ---------------------------------------------------------------------------

@router.get("/api/v1/jobs")
def list_jobs():
    from app.translation.job_store import all_jobs
    jobs = all_jobs()
    return {"total": len(jobs), "jobs": [j.dict() for j in jobs.values()]}